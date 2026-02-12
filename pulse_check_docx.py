from __future__ import annotations

import io
import re
from typing import Dict, Iterable, List

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


MAGENTA = "E20074"
LIGHT_GRAY = "F2F2F2"
BORDER_GRAY = "D9D9D9"
FONT_NAME = "TeleNeo Office"


def load_template(path: str) -> Document:
    return Document(path)


def clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag.endswith('sectPr'):
            continue
        body.remove(child)


def set_default_styles(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0, 0, 0)


def add_title_line(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_section_header(document: Document, text: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _set_cell_borders(cell, color: str = BORDER_GRAY, size: int = 4) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):
        edge_el = tc_borders.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            tc_borders.append(edge_el)
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), str(size))
        edge_el.set(qn("w:color"), color)


def _is_numeric_series(values: Iterable[str]) -> bool:
    pattern = re.compile(r"^[\s\d,\.\-%]+$")
    cleaned = [v for v in values if v not in (None, "")]
    if not cleaned:
        return False
    return all(bool(pattern.match(str(v))) for v in cleaned)


def _apply_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    if color is not None:
        run.font.color.rgb = color


def add_table(document: Document, dataframe: pd.DataFrame, table_type: str, column_widths: List[float]) -> None:
    df = dataframe.copy() if dataframe is not None else pd.DataFrame()
    rows = len(df.index) + 1
    cols = len(df.columns)

    table = document.add_table(rows=rows, cols=cols)
    table.autofit = False

    # Header row
    for col_idx, col_name in enumerate(df.columns):
        cell = table.cell(0, col_idx)
        _apply_cell_text(cell, str(col_name), bold=True, color=RGBColor(255, 255, 255))
        _set_cell_shading(cell, MAGENTA)
        _set_cell_borders(cell)

    # Body rows
    for row_idx in range(len(df.index)):
        row_color = LIGHT_GRAY if row_idx % 2 == 0 else None
        for col_idx in range(cols):
            cell = table.cell(row_idx + 1, col_idx)
            text = "" if pd.isna(df.iat[row_idx, col_idx]) else str(df.iat[row_idx, col_idx])
            _apply_cell_text(cell, text, bold=False, color=RGBColor(0, 0, 0))
            if row_color:
                _set_cell_shading(cell, row_color)
            _set_cell_borders(cell)

    # Alignment (numeric centered, text left)
    for col_idx, col_name in enumerate(df.columns):
        series_values = ["" if pd.isna(v) else str(v) for v in df.iloc[:, col_idx].tolist()]
        is_numeric = _is_numeric_series(series_values)
        for row_idx in range(rows):
            p = table.cell(row_idx, col_idx).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_numeric else WD_ALIGN_PARAGRAPH.LEFT

    # Column widths
    for col_idx, ratio in enumerate(column_widths):
        width = Inches(ratio)
        for row_idx in range(rows):
            table.cell(row_idx, col_idx).width = width


def _content_width_inches(document: Document) -> float:
    section = document.sections[0]
    width = section.page_width - section.left_margin - section.right_margin
    return width / 914400  # EMU to inches


def build_pulse_check_docx(
    template_path: str,
    title_line: str,
    sections: List[Dict[str, object]],
) -> bytes:
    document = load_template(template_path)
    clear_body(document)
    set_default_styles(document)

    add_title_line(document, title_line)

    content_width = _content_width_inches(document)

    for section in sections:
        header = section["header"]
        df = section["data"]
        widths = section["widths"]
        add_section_header(document, header)
        add_table(document, df, header, [content_width * w for w in widths])
        document.add_paragraph("")

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf.getvalue()
